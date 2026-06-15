from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "documentation" / "chapter3.md"
OUTPUT = ROOT / "documentation" / "Ndikum_Chapter_Three_Materials_and_Methods_Revised.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in [(1, 14), (2, 13), (3, 12), (4, 12)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_table(document, rows):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = value
        set_cell_shading(cell, "DCEDE9")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    document.add_paragraph()


def add_paragraph(document, content):
    paragraph = document.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", content)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)
    return paragraph


def add_figure(document, alt, relative_path):
    image_path = ROOT / relative_path
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))
    caption = document.add_paragraph(alt)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(10)


def build():
    document = Document()
    style_document(document)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"-+", cell.replace(" ", "")) for cell in cells):
                    rows.append(cells)
                index += 1
            add_table(document, rows)
            continue
        figure = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", line)
        if figure:
            add_figure(document, figure.group(1), figure.group(2))
            index += 1
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:])
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        elif line.startswith("#### "):
            document.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif re.match(r"^\d+\.\s", line):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(re.sub(r"^\d+\.\s", "", line))
        else:
            add_paragraph(document, line)
        index += 1
    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
